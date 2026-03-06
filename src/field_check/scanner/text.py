"""Text extraction pipeline with ProcessPoolExecutor crash isolation."""

from __future__ import annotations

import logging
import os
from collections.abc import Callable
from concurrent.futures import ProcessPoolExecutor, as_completed
from dataclasses import dataclass, field

from field_check.scanner import FileEntry
from field_check.scanner.inventory import InventoryResult
from field_check.scanner.sampling import SampleResult

logger = logging.getLogger(__name__)

# MIME types that support text extraction
EXTRACTABLE_MIMES: set[str] = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

# Standard metadata fields to check for completeness
METADATA_FIELDS: list[str] = ["title", "author", "creation_date"]

# Image-heavy classification thresholds
CHARS_PER_PAGE_IMAGE_HEAVY = 100
CHARS_PER_PAGE_TEXT_HEAVY = 500
TEXT_SIZE_RATIO_IMAGE_HEAVY = 0.05

# Classification labels
CLASSIFICATION_TEXT_HEAVY = "text_heavy"
CLASSIFICATION_IMAGE_HEAVY = "image_heavy"
CLASSIFICATION_MIXED = "mixed"

# Process pool defaults
DEFAULT_TIMEOUT = 30.0
MAX_WORKERS = 4

# Page count distribution buckets
PAGE_COUNT_BUCKETS: list[tuple[int, float, str]] = [
    (1, 1, "1 page"),
    (2, 5, "2-5 pages"),
    (6, 10, "6-10 pages"),
    (11, 50, "11-50 pages"),
    (51, 100, "51-100 pages"),
    (101, 500, "101-500 pages"),
    (501, float("inf"), ">500 pages"),
]


def _page_count_bucket(count: int) -> str:
    """Map a page count to its distribution bucket label."""
    for low, high, label in PAGE_COUNT_BUCKETS:
        if low <= count <= high:
            return label
    return ">500 pages"


@dataclass
class TextResult:
    """Extraction result for a single file."""

    path: str
    text_length: int = 0
    page_count: int = 0
    chars_per_page: float = 0.0
    text_size_ratio: float = 0.0
    is_scanned: bool = False
    is_mixed_scan: bool = False
    classification: str = ""
    metadata: dict[str, str | None] = field(default_factory=dict)
    error: str | None = None


@dataclass
class TextExtractionResult:
    """Aggregate results from text extraction across sampled files."""

    total_processed: int = 0
    extraction_errors: int = 0
    timeout_errors: int = 0
    scanned_count: int = 0
    native_count: int = 0
    mixed_scan_count: int = 0
    text_heavy_count: int = 0
    image_heavy_count: int = 0
    mixed_content_count: int = 0
    metadata_field_counts: dict[str, int] = field(default_factory=dict)
    metadata_total_checked: int = 0
    page_count_total: int = 0
    page_count_min: int = 0
    page_count_max: int = 0
    page_count_distribution: dict[str, int] = field(default_factory=dict)
    file_results: list[TextResult] = field(default_factory=list)


def _extract_pdf(filepath: str) -> TextResult:
    """Extract text, metadata, and classification from a PDF.

    Single-pass extraction: opens the file once to get text, page count,
    scanned detection (via page.chars), content classification, and metadata.
    """
    import pdfplumber

    result = TextResult(path=filepath)
    file_size = os.path.getsize(filepath)

    try:
        with pdfplumber.open(filepath) as pdf:
            result.page_count = len(pdf.pages)

            # Extract metadata
            meta = pdf.metadata or {}
            result.metadata["title"] = meta.get("Title") or None
            result.metadata["author"] = meta.get("Author") or None
            creation = meta.get("CreationDate")
            result.metadata["creation_date"] = str(creation) if creation else None

            total_chars = 0
            total_text_bytes = 0
            scanned_pages = 0
            native_pages = 0

            for page in pdf.pages:
                # Count char objects for scanned detection
                char_count = len(page.chars) if page.chars else 0
                if char_count == 0:
                    scanned_pages += 1
                else:
                    native_pages += 1
                    total_chars += char_count

                # Extract text
                text = page.extract_text() or ""
                total_text_bytes += len(text.encode("utf-8", errors="replace"))

            result.text_length = total_chars

            # Scanned detection
            if result.page_count > 0:
                if scanned_pages == result.page_count:
                    result.is_scanned = True
                elif scanned_pages > 0 and native_pages > 0:
                    result.is_mixed_scan = True

            # Content classification
            if result.page_count > 0:
                result.chars_per_page = total_chars / result.page_count
                result.text_size_ratio = (
                    total_text_bytes / file_size if file_size > 0 else 0.0
                )

                if (
                    result.is_scanned
                    or result.chars_per_page < CHARS_PER_PAGE_IMAGE_HEAVY
                ):
                    result.classification = CLASSIFICATION_IMAGE_HEAVY
                elif result.chars_per_page > CHARS_PER_PAGE_TEXT_HEAVY:
                    result.classification = CLASSIFICATION_TEXT_HEAVY
                else:
                    # Mixed zone -- check secondary metric
                    if result.text_size_ratio < TEXT_SIZE_RATIO_IMAGE_HEAVY:
                        result.classification = CLASSIFICATION_IMAGE_HEAVY
                    else:
                        result.classification = CLASSIFICATION_MIXED

    except Exception as exc:
        result.error = str(exc)

    return result


def _docx_full_text(doc: object) -> str:
    """Extract all text from a python-docx Document.

    Includes body paragraphs, table cells, headers, and footers.
    """
    parts: list[str] = []

    # Body paragraphs
    for p in doc.paragraphs:  # type: ignore[attr-defined]
        if p.text:
            parts.append(p.text)

    # Table cells
    for table in doc.tables:  # type: ignore[attr-defined]
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)

    # Headers and footers
    for section in doc.sections:  # type: ignore[attr-defined]
        for hf in (section.header, section.footer):
            if hf and hf.is_linked_to_previous is False:
                for p in hf.paragraphs:
                    if p.text:
                        parts.append(p.text)

    return "\n".join(parts)


def _extract_docx(filepath: str) -> TextResult:
    """Extract text and metadata from a DOCX file."""
    from docx import Document

    result = TextResult(path=filepath)
    file_size = os.path.getsize(filepath)

    try:
        doc = Document(filepath)

        # Extract text from paragraphs, tables, headers, footers
        text = _docx_full_text(doc)
        text_bytes = len(text.encode("utf-8", errors="replace"))
        result.text_length = len(text)
        result.text_size_ratio = text_bytes / file_size if file_size > 0 else 0.0

        # Extract metadata
        props = doc.core_properties
        result.metadata["title"] = props.title if props.title else None
        result.metadata["author"] = props.author if props.author else None
        result.metadata["creation_date"] = (
            props.created.isoformat() if props.created else None
        )

        # DOCX is always text-based, never scanned
        result.classification = CLASSIFICATION_TEXT_HEAVY

    except Exception as exc:
        result.error = str(exc)

    return result


def _extract_single(filepath: str, mime_type: str) -> TextResult:
    """Worker function for ProcessPoolExecutor.

    Must be top-level for pickling. Dispatches to type-specific extractors.
    """
    if mime_type == "application/pdf":
        return _extract_pdf(filepath)
    elif (
        mime_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        return _extract_docx(filepath)
    else:
        return TextResult(path=filepath, error=f"Unsupported type: {mime_type}")


def extract_text(
    sample: SampleResult,
    inventory: InventoryResult,
    max_workers: int | None = None,
    timeout: float = DEFAULT_TIMEOUT,
    progress_callback: Callable[[int, int], None] | None = None,
) -> TextExtractionResult:
    """Extract text from sampled PDF/DOCX files using process pool.

    Uses ProcessPoolExecutor for crash isolation -- if a worker segfaults
    on a malformed file, the pool spawns a new worker automatically.

    Args:
        sample: Sampling result with selected files.
        inventory: Inventory with per-file MIME type mapping.
        max_workers: Max worker processes (default: min(4, cpu_count)).
        timeout: Per-file timeout in seconds.
        progress_callback: Called with (current, total) for progress display.

    Returns:
        Aggregated text extraction results.
    """
    result = TextExtractionResult()

    # Filter to extractable types only
    extractable: list[tuple[FileEntry, str]] = []
    for entry in sample.selected_files:
        mime = inventory.file_types.get(entry.path, "application/octet-stream")
        if mime in EXTRACTABLE_MIMES:
            extractable.append((entry, mime))

    if not extractable:
        return result

    total = len(extractable)
    workers = max_workers or min(MAX_WORKERS, os.cpu_count() or 1)

    with ProcessPoolExecutor(max_workers=workers) as pool:
        future_to_info: dict = {}
        for entry, mime in extractable:
            future = pool.submit(_extract_single, str(entry.path), mime)
            future_to_info[future] = (entry, mime)

        for completed, future in enumerate(as_completed(future_to_info), 1):
            entry, mime = future_to_info[future]
            try:
                text_result = future.result(timeout=timeout)
            except TimeoutError:
                text_result = TextResult(
                    path=str(entry.path), error="Extraction timed out"
                )
                result.timeout_errors += 1
            except Exception as exc:
                text_result = TextResult(path=str(entry.path), error=str(exc))

            result.file_results.append(text_result)
            result.total_processed += 1

            if text_result.error:
                result.extraction_errors += 1
            else:
                # Scanned detection counts (PDF-only concept)
                if mime == "application/pdf":
                    if text_result.is_scanned:
                        result.scanned_count += 1
                    elif text_result.is_mixed_scan:
                        result.mixed_scan_count += 1
                    else:
                        result.native_count += 1

                # Content classification counts
                if text_result.classification == CLASSIFICATION_TEXT_HEAVY:
                    result.text_heavy_count += 1
                elif text_result.classification == CLASSIFICATION_IMAGE_HEAVY:
                    result.image_heavy_count += 1
                elif text_result.classification == CLASSIFICATION_MIXED:
                    result.mixed_content_count += 1

                # Metadata completeness
                result.metadata_total_checked += 1
                for mf in METADATA_FIELDS:
                    value = text_result.metadata.get(mf)
                    if value:
                        result.metadata_field_counts[mf] = (
                            result.metadata_field_counts.get(mf, 0) + 1
                        )

                # Page count distribution (PDF only)
                if text_result.page_count > 0:
                    result.page_count_total += text_result.page_count
                    if (
                        result.page_count_min == 0
                        or text_result.page_count < result.page_count_min
                    ):
                        result.page_count_min = text_result.page_count
                    if text_result.page_count > result.page_count_max:
                        result.page_count_max = text_result.page_count
                    bucket = _page_count_bucket(text_result.page_count)
                    result.page_count_distribution[bucket] = (
                        result.page_count_distribution.get(bucket, 0) + 1
                    )

            if progress_callback is not None:
                progress_callback(completed, total)

    return result


# ---------------------------------------------------------------------------
# Shared text cache for downstream analysis (PII, language, encoding)
# ---------------------------------------------------------------------------

# MIME types for plain text content extraction
PLAIN_TEXT_MIMES: set[str] = {
    "text/plain",
    "text/csv",
    "text/json",
    "text/xml",
    "application/json",
    "application/xml",
}

# Combined set: all types we can extract text from for the cache
CACHE_EXTRACTABLE_MIMES: set[str] = EXTRACTABLE_MIMES | PLAIN_TEXT_MIMES

# Max bytes to read from plain text files
_MAX_TEXT_READ = 1_000_000


@dataclass
class TextCacheResult:
    """Result of shared text cache extraction."""

    text_cache: dict[str, str] = field(default_factory=dict)
    encoding_map: dict[str, tuple[str, float]] = field(default_factory=dict)
    total_extracted: int = 0
    extraction_errors: int = 0


def _extract_text_for_cache(
    filepath: str, mime_type: str
) -> tuple[str, str | None, float, str | None]:
    """Worker function to extract text content from any supported file.

    Must be top-level for pickling (ProcessPoolExecutor).

    Returns:
        Tuple of (text, encoding_name, encoding_confidence, error).
        encoding_name is only set for plain text files.
    """
    try:
        if mime_type == "application/pdf":
            import pdfplumber

            with pdfplumber.open(filepath) as pdf:
                text = "\n".join(
                    page.extract_text() or "" for page in pdf.pages
                )
            return (text, None, 0.0, None)

        if (
            mime_type
            == "application/vnd.openxmlformats-officedocument"
               ".wordprocessingml.document"
        ):
            from docx import Document

            doc = Document(filepath)
            text = _docx_full_text(doc)
            return (text, None, 0.0, None)

        # Plain text types — read bytes, detect encoding
        from charset_normalizer import from_bytes

        with open(filepath, "rb") as f:
            raw = f.read(_MAX_TEXT_READ)

        result = from_bytes(raw).best()
        if result:
            return (str(result), result.encoding, 1.0 - result.chaos, None)
        return (
            raw.decode("utf-8", errors="replace"),
            "utf-8",
            0.0,
            None,
        )
    except Exception as exc:
        return ("", None, 0.0, str(exc))


def build_text_cache(
    sample: SampleResult,
    inventory: InventoryResult,
    max_workers: int | None = None,
    timeout: float = DEFAULT_TIMEOUT,
    progress_callback: Callable[[int, int], None] | None = None,
) -> TextCacheResult:
    """Build shared text cache for downstream analysis (PII, language, encoding).

    Uses ProcessPoolExecutor for crash isolation. Extracts text from:
    - PDFs via pdfplumber
    - DOCXes via python-docx
    - Plain text files via charset-normalizer (also captures encoding)

    Args:
        sample: Sampling result with selected files.
        inventory: Inventory with per-file MIME type mapping.
        max_workers: Max worker processes.
        timeout: Per-file timeout in seconds.
        progress_callback: Called with (current, total).

    Returns:
        TextCacheResult with text_cache dict and encoding_map.
    """
    cache_result = TextCacheResult()

    # Filter to cache-extractable types
    extractable: list[tuple[FileEntry, str]] = []
    for entry in sample.selected_files:
        mime = inventory.file_types.get(entry.path, "application/octet-stream")
        if mime in CACHE_EXTRACTABLE_MIMES:
            extractable.append((entry, mime))

    if not extractable:
        return cache_result

    total = len(extractable)
    workers = max_workers or min(MAX_WORKERS, os.cpu_count() or 1)

    with ProcessPoolExecutor(max_workers=workers) as pool:
        future_to_entry: dict = {}
        for entry, mime in extractable:
            future = pool.submit(
                _extract_text_for_cache, str(entry.path), mime
            )
            future_to_entry[future] = entry

        for completed, future in enumerate(as_completed(future_to_entry), 1):
            entry = future_to_entry[future]
            path_str = str(entry.path)
            try:
                text, enc_name, enc_conf, error = future.result(
                    timeout=timeout
                )
            except TimeoutError:
                cache_result.extraction_errors += 1
                cache_result.total_extracted += 1
                if progress_callback is not None:
                    progress_callback(completed, total)
                continue
            except Exception:
                cache_result.extraction_errors += 1
                cache_result.total_extracted += 1
                if progress_callback is not None:
                    progress_callback(completed, total)
                continue

            cache_result.total_extracted += 1

            if error:
                cache_result.extraction_errors += 1
            else:
                if text:
                    cache_result.text_cache[path_str] = text
                if enc_name:
                    cache_result.encoding_map[path_str] = (
                        enc_name,
                        enc_conf,
                    )

            if progress_callback is not None:
                progress_callback(completed, total)

    return cache_result
